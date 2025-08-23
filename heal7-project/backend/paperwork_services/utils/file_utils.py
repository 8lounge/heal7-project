#!/usr/bin/env python3
"""
파일 처리 유틸리티
"""

import os
import base64
import mimetypes
from pathlib import Path
from typing import Dict, Any, List, Optional
import tempfile
import logging

# 문서 처리 라이브러리
try:
    import PyPDF2
    from docx import Document
except ImportError as e:
    logging.warning(f"문서 처리 라이브러리 임포트 실패: {e}")
    PyPDF2 = None
    Document = None

class FileUtils:
    """파일 처리 유틸리티 클래스"""
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        """파일 확장자 반환"""
        return Path(filename).suffix.lower()
    
    @staticmethod
    def is_image_file(filename: str) -> bool:
        """이미지 파일인지 확인"""
        image_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
        return FileUtils.get_file_extension(filename) in image_extensions
    
    @staticmethod
    def is_document_file(filename: str) -> bool:
        """문서 파일인지 확인"""
        doc_extensions = ['.pdf', '.hwp', '.hwpx', '.doc', '.docx']
        return FileUtils.get_file_extension(filename) in doc_extensions
    
    @staticmethod
    def get_mime_type(filename: str) -> str:
        """파일의 MIME 타입 반환"""
        mime_type, _ = mimetypes.guess_type(filename)
        return mime_type or 'application/octet-stream'
    
    @staticmethod
    def encode_file_to_base64(file_path: str) -> str:
        """파일을 Base64로 인코딩"""
        try:
            with open(file_path, 'rb') as f:
                return base64.b64encode(f.read()).decode('utf-8')
        except Exception as e:
            raise ValueError(f"파일 인코딩 실패: {e}")
    
    @staticmethod
    def decode_base64_to_file(base64_data: str, output_path: str) -> bool:
        """Base64 데이터를 파일로 저장"""
        try:
            decoded_data = base64.b64decode(base64_data)
            with open(output_path, 'wb') as f:
                f.write(decoded_data)
            return True
        except Exception as e:
            print(f"파일 디코딩 실패: {e}")
            return False
    
    @staticmethod
    def create_temp_file(suffix: str = '', prefix: str = 'paperwork_') -> str:
        """임시 파일 생성"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix, prefix=prefix) as temp_file:
            return temp_file.name
    
    @staticmethod
    def cleanup_temp_file(file_path: str) -> bool:
        """임시 파일 정리"""
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
                return True
            return False
        except Exception as e:
            print(f"임시 파일 정리 실패: {e}")
            return False
    
    @staticmethod
    def get_file_size(file_path: str) -> int:
        """파일 크기 반환 (bytes)"""
        try:
            return os.path.getsize(file_path)
        except Exception:
            return 0
    
    @staticmethod
    def is_file_size_valid(file_path: str, max_size: int = 50 * 1024 * 1024) -> bool:
        """파일 크기가 유효한지 확인 (기본 50MB)"""
        file_size = FileUtils.get_file_size(file_path)
        return 0 < file_size <= max_size
    
    @staticmethod
    def validate_file(file_path: str, allowed_extensions: List[str] = None, max_size: int = None) -> Dict[str, Any]:
        """파일 유효성 검사"""
        if not os.path.exists(file_path):
            return {'valid': False, 'error': '파일이 존재하지 않습니다'}
        
        filename = os.path.basename(file_path)
        file_ext = FileUtils.get_file_extension(filename)
        
        # 확장자 검사
        if allowed_extensions and file_ext not in allowed_extensions:
            return {
                'valid': False, 
                'error': f'지원하지 않는 파일 형식: {file_ext}',
                'allowed_extensions': allowed_extensions
            }
        
        # 파일 크기 검사
        if max_size and not FileUtils.is_file_size_valid(file_path, max_size):
            file_size = FileUtils.get_file_size(file_path)
            return {
                'valid': False,
                'error': f'파일 크기 초과: {file_size} bytes (최대: {max_size} bytes)',
                'file_size': file_size,
                'max_size': max_size
            }
        
        return {
            'valid': True,
            'filename': filename,
            'extension': file_ext,
            'size': FileUtils.get_file_size(file_path),
            'mime_type': FileUtils.get_mime_type(filename)
        }
    
    @staticmethod
    def extract_text_from_pdf(pdf_file_path: str) -> str:
        """PDF 파일에서 텍스트 직접 추출"""
        try:
            if not PyPDF2:
                return f"[PDF 처리 라이브러리 없음] {os.path.basename(pdf_file_path)}"
            
            text_content = []
            
            with open(pdf_file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"=== 페이지 {page_num} ===\n{page_text}\n")
                    except Exception as e:
                        text_content.append(f"=== 페이지 {page_num} (오류) ===\n텍스트 추출 실패: {str(e)}\n")
            
            return '\n'.join(text_content) if text_content else "[PDF에서 텍스트를 찾을 수 없습니다]"
            
        except Exception as e:
            return f"[PDF 처리 오류: {str(e)}] {os.path.basename(pdf_file_path)}"
    
    @staticmethod
    def extract_text_from_docx(docx_file_path: str) -> str:
        """DOCX 파일에서 텍스트 직접 추출"""
        try:
            if not Document:
                return f"[DOCX 처리 라이브러리 없음] {os.path.basename(docx_file_path)}"
            
            doc = Document(docx_file_path)
            text_content = []
            
            # 문서의 모든 단락 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 표 내용 추출
            for table in doc.tables:
                text_content.append("\n[표 시작]")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
                text_content.append("[표 끝]\n")
            
            return '\n'.join(text_content) if text_content else "[DOCX에서 텍스트를 찾을 수 없습니다]"
            
        except Exception as e:
            return f"[DOCX 처리 오류: {str(e)}] {os.path.basename(docx_file_path)}"
    
    @staticmethod
    def extract_text_from_doc(doc_file_path: str) -> str:
        """DOC 파일에서 텍스트 추출 (제한적 지원)"""
        # python-docx는 .doc 파일을 직접 지원하지 않음
        filename = os.path.basename(doc_file_path)
        return f"[DOC 파일 변환 필요: {filename}]\n\n.doc 파일은 .docx로 변환 후 처리하거나 LibreOffice 등의 도구가 필요합니다."
    
    @staticmethod
    def extract_text_from_hwp(hwp_file_path: str) -> str:
        """HWP 파일에서 텍스트 추출 (OCR 방식 사용 권장)"""
        filename = os.path.basename(hwp_file_path)
        return f"[HWP 파일: {filename}]\n\nHWP 파일은 이미지 변환 후 OCR 처리를 권장합니다. 직접 텍스트 추출을 위해서는 pyhwp 라이브러리가 필요합니다."
    
    @staticmethod
    def extract_text_from_document(file_path: str) -> Dict[str, Any]:
        """
        문서 타입에 따라 최적의 텍스트 추출 방식 선택
        - PDF, DOCX: 직접 텍스트 추출
        - HWP, HWPX: OCR 방식 권장 (이 메서드에서는 placeholder 반환)
        - DOC: 변환 필요 안내
        """
        if not os.path.exists(file_path):
            return {
                'success': False,
                'error': '파일이 존재하지 않습니다',
                'text': '',
                'method': 'none'
            }
        
        file_ext = FileUtils.get_file_extension(os.path.basename(file_path)).lower()
        
        try:
            if file_ext == '.pdf':
                text = FileUtils.extract_text_from_pdf(file_path)
                return {
                    'success': True,
                    'text': text,
                    'method': 'direct_pdf',
                    'file_type': 'pdf',
                    'recommended': True
                }
            
            elif file_ext == '.docx':
                text = FileUtils.extract_text_from_docx(file_path)
                return {
                    'success': True,
                    'text': text,
                    'method': 'direct_docx',
                    'file_type': 'docx',
                    'recommended': True
                }
            
            elif file_ext == '.doc':
                text = FileUtils.extract_text_from_doc(file_path)
                return {
                    'success': True,
                    'text': text,
                    'method': 'conversion_required',
                    'file_type': 'doc',
                    'recommended': False,
                    'note': 'DOC 파일은 DOCX로 변환 후 처리 권장'
                }
            
            elif file_ext in ['.hwp', '.hwpx']:
                text = FileUtils.extract_text_from_hwp(file_path)
                return {
                    'success': True,
                    'text': text,
                    'method': 'ocr_recommended',
                    'file_type': file_ext.replace('.', ''),
                    'recommended': False,
                    'note': 'HWP 파일은 이미지 변환 후 OCR 처리가 더 정확합니다'
                }
            
            else:
                return {
                    'success': False,
                    'error': f'지원하지 않는 파일 형식: {file_ext}',
                    'text': '',
                    'method': 'unsupported'
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f'문서 처리 중 오류: {str(e)}',
                'text': '',
                'method': 'error'
            }
    
    @staticmethod  
    def should_use_ocr(file_path: str) -> bool:
        """파일이 OCR 처리를 사용해야 하는지 판단"""
        file_ext = FileUtils.get_file_extension(os.path.basename(file_path)).lower()
        return file_ext in ['.hwp', '.hwpx', '.png', '.jpg', '.jpeg', '.tiff', '.bmp']
    
    @staticmethod
    def get_file_info(file_path: str) -> Dict[str, Any]:
        """파일 정보 반환"""
        if not os.path.exists(file_path):
            return {'exists': False}
        
        stat_info = os.stat(file_path)
        filename = os.path.basename(file_path)
        
        return {
            'exists': True,
            'filename': filename,
            'extension': FileUtils.get_file_extension(filename),
            'size': stat_info.st_size,
            'mime_type': FileUtils.get_mime_type(filename),
            'is_image': FileUtils.is_image_file(filename),
            'is_document': FileUtils.is_document_file(filename),
            'created_time': stat_info.st_ctime,
            'modified_time': stat_info.st_mtime
        }