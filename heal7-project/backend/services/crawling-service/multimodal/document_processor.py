#!/usr/bin/env python3
"""
📄 문서 처리기
PDF, DOCX, XLSX 등 다양한 문서 형식 처리 및 AI 분석

Features:
- PDF 텍스트 추출 및 이미지 변환
- Word 문서 처리
- Excel 파일 분석
- 이미지 문서 OCR
- AI 기반 문서 분석

Author: HEAL7 Development Team
Version: 1.0.0
Date: 2025-08-30
"""

import os
import asyncio
import logging
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
from dataclasses import dataclass
from enum import Enum

# 문서 처리 라이브러리
try:
    import PyPDF2
    from pdf2image import convert_from_path
except ImportError:
    PyPDF2 = None
    convert_from_path = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

from PIL import Image
import io

from .ai_analyzer import MultimodalAnalyzer, AIModel


logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """문서 타입"""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    PPTX = "pptx"
    HWP = "hwp"
    IMAGE = "image"
    TEXT = "text"
    UNKNOWN = "unknown"


@dataclass
class DocumentAnalysisResult:
    """문서 분석 결과"""
    success: bool
    document_type: DocumentType
    file_path: str
    
    # 추출된 콘텐츠
    text_content: List[str] = None
    tables: List[Dict] = None
    images: List[Dict] = None
    metadata: Dict[str, Any] = None
    
    # AI 분석 결과
    ai_summary: Optional[str] = None
    key_points: List[str] = None
    
    # 에러 정보
    error: Optional[str] = None
    processing_time: float = 0.0
    
    def __post_init__(self):
        if self.text_content is None:
            self.text_content = []
        if self.tables is None:
            self.tables = []
        if self.images is None:
            self.images = []
        if self.metadata is None:
            self.metadata = {}
        if self.key_points is None:
            self.key_points = []


class DocumentProcessor:
    """📄 통합 문서 처리기"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.DocumentProcessor")
        self.ai_analyzer = MultimodalAnalyzer()
        
        # 지원하는 문서 형식
        self.supported_formats = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOCX,  # 레거시 지원
            '.xlsx': DocumentType.XLSX,
            '.xls': DocumentType.XLSX,
            '.pptx': DocumentType.PPTX,
            '.hwp': DocumentType.HWP,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE,
            '.png': DocumentType.IMAGE,
            '.gif': DocumentType.IMAGE,
            '.bmp': DocumentType.IMAGE,
            '.webp': DocumentType.IMAGE,
            '.txt': DocumentType.TEXT,
            '.md': DocumentType.TEXT,
            '.log': DocumentType.TEXT
        }
        
        # 처리기 매핑
        self.processors = {
            DocumentType.PDF: self._process_pdf,
            DocumentType.DOCX: self._process_docx,
            DocumentType.XLSX: self._process_xlsx,
            DocumentType.IMAGE: self._process_image_document,
            DocumentType.TEXT: self._process_text,
        }
    
    async def initialize(self):
        """문서 처리기 초기화"""
        self.logger.info("📄 문서 처리기 초기화")
        await self.ai_analyzer.initialize()
        self.logger.info("✅ 문서 처리기 초기화 완료")
    
    async def process_document(
        self, 
        file_path: str, 
        include_ai_analysis: bool = True,
        extract_images: bool = True
    ) -> DocumentAnalysisResult:
        """문서 처리 메인 함수"""
        
        start_time = asyncio.get_event_loop().time()
        
        try:
            # 파일 존재 확인
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
            
            # 문서 타입 결정
            doc_type = self._determine_document_type(file_path)
            
            self.logger.info(f"📄 문서 처리 시작: {doc_type.value} - {file_path}")
            
            # 기본 결과 객체 생성
            result = DocumentAnalysisResult(
                success=False,
                document_type=doc_type,
                file_path=file_path
            )
            
            # 문서 타입별 처리
            if doc_type in self.processors:
                await self.processors[doc_type](file_path, result, extract_images)
            else:
                # 지원하지 않는 형식은 이미지 변환 후 처리
                await self._fallback_image_processing(file_path, result)
            
            # AI 분석 (선택사항)
            if include_ai_analysis and result.success:
                await self._add_ai_analysis(result)
            
            result.success = True
            result.processing_time = asyncio.get_event_loop().time() - start_time
            
            self.logger.info(f"✅ 문서 처리 완료: {result.processing_time:.2f}초")
            return result
            
        except Exception as e:
            error_msg = f"문서 처리 실패: {e}"
            self.logger.error(error_msg)
            
            return DocumentAnalysisResult(
                success=False,
                document_type=DocumentType.UNKNOWN,
                file_path=file_path,
                error=error_msg,
                processing_time=asyncio.get_event_loop().time() - start_time
            )
    
    def _determine_document_type(self, file_path: str) -> DocumentType:
        """파일 확장자로 문서 타입 결정"""
        extension = Path(file_path).suffix.lower()
        return self.supported_formats.get(extension, DocumentType.UNKNOWN)
    
    async def _process_pdf(self, file_path: str, result: DocumentAnalysisResult, extract_images: bool):
        """PDF 문서 처리"""
        
        if not PyPDF2:
            raise ImportError("PyPDF2 라이브러리가 필요합니다")
        
        # 1. 텍스트 추출
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            result.metadata = {
                'pages': len(pdf_reader.pages),
                'title': pdf_reader.metadata.get('/Title', ''),
                'author': pdf_reader.metadata.get('/Author', ''),
                'subject': pdf_reader.metadata.get('/Subject', ''),
                'creator': pdf_reader.metadata.get('/Creator', '')
            }
            
            # 각 페이지의 텍스트 추출
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        result.text_content.append(text)
                except:
                    # 텍스트 추출 실패한 페이지는 건너뛰기
                    continue
        
        # 2. 이미지로 변환하여 테이블/이미지 추출
        if extract_images and convert_from_path:
            try:
                images = convert_from_path(file_path, dpi=150)
                
                for page_num, image in enumerate(images):
                    # 이미지를 바이트로 변환
                    img_bytes = self._pil_to_bytes(image)
                    
                    # 테이블이 있는지 AI로 확인
                    if await self._has_table_in_image(img_bytes):
                        table_data = await self.ai_analyzer.extract_table_from_image(img_bytes)
                        
                        if table_data['success'] and table_data.get('parsed_tables'):
                            result.tables.extend(table_data['parsed_tables']['tables'])
                    
                    # 이미지 정보 저장
                    result.images.append({
                        'page': page_num + 1,
                        'size': image.size,
                        'has_table': len([t for t in result.tables if t.get('page') == page_num + 1]) > 0
                    })
                    
            except Exception as e:
                self.logger.warning(f"PDF 이미지 처리 실패: {e}")
    
    async def _process_docx(self, file_path: str, result: DocumentAnalysisResult, extract_images: bool):
        """Word 문서 처리"""
        
        if not Document:
            raise ImportError("python-docx 라이브러리가 필요합니다")
        
        doc = Document(file_path)
        
        # 문서 속성
        result.metadata = {
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
            'subject': doc.core_properties.subject or '',
            'created': str(doc.core_properties.created) if doc.core_properties.created else '',
            'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
        }
        
        # 단락별 텍스트 추출
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                result.text_content.append(paragraph.text)
        
        # 표 추출
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'table_index': table_idx + 1,
                'headers': [],
                'rows': []
            }
            
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                
                if row_idx == 0:
                    table_data['headers'] = row_data
                else:
                    table_data['rows'].append(row_data)
            
            result.tables.append(table_data)
    
    async def _process_xlsx(self, file_path: str, result: DocumentAnalysisResult, extract_images: bool):
        """Excel 파일 처리"""
        
        if not load_workbook:
            raise ImportError("openpyxl 라이브러리가 필요합니다")
        
        workbook = load_workbook(file_path, data_only=True)
        
        result.metadata = {
            'worksheets': workbook.sheetnames,
            'active_sheet': workbook.active.title if workbook.active else ''
        }
        
        # 각 워크시트 처리
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            
            # 데이터가 있는 영역 찾기
            if worksheet.max_row > 1 and worksheet.max_column > 1:
                # 첫 번째 행을 헤더로 가정
                headers = []
                for col in range(1, worksheet.max_column + 1):
                    cell_value = worksheet.cell(1, col).value
                    headers.append(str(cell_value) if cell_value else f'Column{col}')
                
                # 데이터 행 추출
                rows = []
                for row in range(2, min(worksheet.max_row + 1, 1000)):  # 최대 1000행
                    row_data = []
                    has_data = False
                    
                    for col in range(1, worksheet.max_column + 1):
                        cell_value = worksheet.cell(row, col).value
                        row_data.append(cell_value)
                        if cell_value is not None:
                            has_data = True
                    
                    if has_data:
                        rows.append(row_data)
                    elif len(rows) > 0:  # 빈 행이 나타나면 중단
                        break
                
                if rows:
                    result.tables.append({
                        'sheet_name': sheet_name,
                        'headers': headers,
                        'rows': rows
                    })
            
            # 텍스트 데이터도 추가
            text_content = f"=== {sheet_name} ===\n"
            for row in worksheet.iter_rows(max_row=min(100, worksheet.max_row), values_only=True):
                row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                if row_text.strip():
                    text_content += row_text + '\n'
            
            result.text_content.append(text_content)
    
    async def _process_image_document(self, file_path: str, result: DocumentAnalysisResult, extract_images: bool):
        """이미지 문서 처리 (OCR)"""
        
        with open(file_path, 'rb') as f:
            image_bytes = f.read()
        
        # 이미지 메타데이터
        try:
            image = Image.open(file_path)
            result.metadata = {
                'size': image.size,
                'format': image.format,
                'mode': image.mode
            }
        except:
            pass
        
        # OCR 텍스트 추출
        ocr_result = await self.ai_analyzer.ocr_extract_text(image_bytes)
        if ocr_result['success']:
            result.text_content.append(ocr_result['content'])
        
        # 테이블 추출 시도
        table_result = await self.ai_analyzer.extract_table_from_image(image_bytes)
        if table_result['success'] and table_result.get('parsed_tables'):
            result.tables.extend(table_result['parsed_tables']['tables'])
    
    async def _process_text(self, file_path: str, result: DocumentAnalysisResult, extract_images: bool):
        """텍스트 파일 처리"""
        
        try:
            # 텍스트 파일 읽기 (다양한 인코딩 시도)
            encodings = ['utf-8', 'cp949', 'latin-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise Exception("지원되는 인코딩으로 파일을 읽을 수 없습니다")
            
            # 텍스트 내용 저장
            result.text_content = [content]
            
            # 파일 메타데이터
            import os
            file_stat = os.stat(file_path)
            result.metadata = {
                'size_bytes': file_stat.st_size,
                'lines': len(content.split('\n')),
                'characters': len(content),
                'encoding': encoding
            }
            
        except Exception as e:
            raise Exception(f"텍스트 파일 처리 실패: {e}")
    
    async def _fallback_image_processing(self, file_path: str, result: DocumentAnalysisResult):
        """지원하지 않는 형식을 이미지로 변환 후 처리"""
        try:
            # 이미지 변환 시도
            image = Image.open(file_path)
            image = image.convert('RGB')  # JPEG 호환성
            
            # 임시 JPEG 파일 생성
            temp_path = file_path + '_temp.jpg'
            image.save(temp_path, 'JPEG')
            
            # 이미지 문서로 처리
            await self._process_image_document(temp_path, result, True)
            
            # 임시 파일 삭제
            os.unlink(temp_path)
            
        except Exception as e:
            raise Exception(f"이미지 변환 실패: {e}")
    
    async def _has_table_in_image(self, image_bytes: bytes) -> bool:
        """이미지에 테이블이 있는지 확인"""
        try:
            prompt = "이 이미지에 표(table)가 있습니까? 'Yes' 또는 'No'로만 답변하세요."
            result = await self.ai_analyzer.analyze_image(image_bytes, prompt)
            
            if result['success']:
                return 'yes' in result['content'].lower()
            
        except:
            pass
        
        return False
    
    async def _add_ai_analysis(self, result: DocumentAnalysisResult):
        """AI 기반 문서 분석 추가"""
        try:
            # 텍스트 내용 합치기
            full_text = '\n'.join(result.text_content)
            
            if not full_text.strip():
                return
            
            # 문서 요약
            summary_prompt = f"""
            다음 문서를 요약하고 주요 포인트를 추출하세요:
            
            {full_text[:4000]}  # 처음 4000자만 사용
            
            다음 형식으로 답변하세요:
            
            ## 요약
            [문서의 주요 내용을 2-3문장으로 요약]
            
            ## 주요 포인트
            - [핵심 포인트 1]
            - [핵심 포인트 2]
            - [핵심 포인트 3]
            """
            
            analysis_result = await self.ai_analyzer.analyze_image(
                self._text_to_image(summary_prompt), 
                "텍스트 이미지를 분석하여 요약하고 주요 포인트를 추출하세요."
            )
            
            if analysis_result['success']:
                content = analysis_result['content']
                
                # 요약과 주요 포인트 분리
                if '## 요약' in content and '## 주요 포인트' in content:
                    parts = content.split('## 주요 포인트')
                    summary_part = parts[0].replace('## 요약', '').strip()
                    points_part = parts[1].strip()
                    
                    result.ai_summary = summary_part
                    
                    # 주요 포인트 추출
                    for line in points_part.split('\n'):
                        if line.strip().startswith('-'):
                            result.key_points.append(line.strip()[1:].strip())
                
        except Exception as e:
            self.logger.warning(f"AI 분석 실패: {e}")
    
    def _pil_to_bytes(self, image: Image.Image) -> bytes:
        """PIL Image를 바이트로 변환"""
        output = io.BytesIO()
        image.save(output, format='JPEG', quality=90)
        return output.getvalue()
    
    def _text_to_image(self, text: str) -> bytes:
        """텍스트를 이미지로 변환 (AI 분석용)"""
        from PIL import Image, ImageDraw, ImageFont
        
        # 이미지 크기 계산
        lines = text.split('\n')
        max_width = max(len(line) for line in lines) * 10
        height = len(lines) * 20 + 40
        
        # 이미지 생성
        image = Image.new('RGB', (min(max_width, 800), min(height, 1000)), 'white')
        draw = ImageDraw.Draw(image)
        
        # 폰트 설정 (기본 폰트 사용)
        try:
            font = ImageFont.load_default()
        except:
            font = None
        
        # 텍스트 그리기
        y_position = 20
        for line in lines:
            draw.text((20, y_position), line, fill='black', font=font)
            y_position += 20
        
        # 바이트로 변환
        output = io.BytesIO()
        image.save(output, format='JPEG')
        return output.getvalue()
    
    def get_supported_formats(self) -> List[str]:
        """지원하는 파일 형식 목록"""
        return list(self.supported_formats.keys())


# 유틸리티 함수들

async def quick_document_analysis(file_path: str, include_ai: bool = True) -> DocumentAnalysisResult:
    """빠른 문서 분석"""
    processor = DocumentProcessor()
    await processor.initialize()
    
    return await processor.process_document(file_path, include_ai)


async def extract_tables_from_document(file_path: str) -> List[Dict]:
    """문서에서 테이블만 추출"""
    result = await quick_document_analysis(file_path, include_ai=False)
    
    if result.success:
        return result.tables
    else:
        return []


async def batch_document_processing(file_paths: List[str], max_concurrent: int = 3) -> List[DocumentAnalysisResult]:
    """배치 문서 처리"""
    processor = DocumentProcessor()
    await processor.initialize()
    
    semaphore = asyncio.Semaphore(max_concurrent)
    
    async def process_single(file_path: str):
        async with semaphore:
            return await processor.process_document(file_path)
    
    tasks = [process_single(path) for path in file_paths]
    return await asyncio.gather(*tasks, return_exceptions=True)